from docx import Document
from docx2pdf import convert
import os
import fitz


def find_and_replace(file_path, replacements, save_path):
    document = Document(file_path)

    # Handle paragraphs
    for paragraph in document.paragraphs:
        replace_in_paragraph(paragraph, replacements)

    # Handle tables
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_in_paragraph(cell, replacements)

    # Handle headers and footers
    for section in document.sections:
        for header in section.header.paragraphs:
            replace_in_paragraph(header, replacements)
        for footer in section.footer.paragraphs:
            replace_in_paragraph(footer, replacements)

    # Save the updated content to the new location
    document.save(save_path)


def replace_in_paragraph(element, replacements):
    # Check if the element is a table cell
    if hasattr(element, "paragraphs"):
        for paragraph in element.paragraphs:
            replace_text_in_paragraph(paragraph, replacements)
    # Otherwise, assume it's a paragraph
    else:
        replace_text_in_paragraph(element, replacements)


def replace_text_in_paragraph(paragraph, replacements):
    for run in paragraph.runs:
        for find_text, replace_text in replacements.items():
            run.text = run.text.replace(find_text, replace_text)


# Funciton to convert docx to a pdf and save it to output path
def convert_docx_to_pdf(input_path, output_path):
    # convert to pdf
    convert(input_path, "temp_pdf.pdf")

    # remove second blank page
    file_handle = fitz.open("temp_pdf.pdf")
    file_handle.delete_page(1)
    file_handle.save(output_path)
